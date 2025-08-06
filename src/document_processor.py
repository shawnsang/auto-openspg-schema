#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理模块

支持多种文档格式的文本提取和分块处理。
"""

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import chardet

# 导入日志
from .logger import logger

# 文档处理库
try:
    import PyPDF2
    logger.success("PyPDF2 库加载成功")
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 库未安装，PDF 处理功能不可用")

try:
    from docx import Document
    logger.success("python-docx 库加载成功")
except ImportError:
    Document = None
    logger.warning("python-docx 库未安装，DOCX 处理功能不可用")

class DocumentProcessor:
    """文档处理器，支持多种格式的文档读取和分块"""
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 200):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 文档分块大小
            chunk_overlap: 分块重叠大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"文档处理器初始化完成 - 分块大小: {chunk_size}, 重叠大小: {chunk_overlap}")
    
    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理单个文档，提取文本并分块
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            List[Dict]: 分块结果列表，每个元素包含 text, metadata 等信息
        """
        logger.info(f"开始处理文档: {file_path}")
        
        try:
            # 根据文件扩展名选择处理方法
            file_ext = os.path.splitext(file_path)[1].lower()
            logger.debug(f"检测到文件格式: {file_ext}")
            
            if file_ext == '.pdf':
                logger.info("使用 PDF 提取器处理文档")
                text = self._extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                logger.info("使用 DOCX 提取器处理文档")
                text = self._extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                logger.info("使用 TXT 提取器处理文档")
                text = self._extract_text_from_txt(file_path)
            else:
                logger.error(f"不支持的文件格式: {file_ext}")
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            logger.success(f"文本提取完成，原始长度: {len(text)} 字符")
            
            # 清理和预处理文本
            logger.debug("开始清理文本")
            text = self._clean_text(text)
            logger.info(f"文本清理完成，清理后长度: {len(text)} 字符")
            
            # 分块处理
            logger.debug("开始文本分块")
            chunks = self._split_text_into_chunks(text)
            logger.success(f"文本分块完成，共生成 {len(chunks)} 个分块")
            
            # 返回结构化的分块数据
            results = []
            for i, chunk in enumerate(chunks):
                results.append({
                    'content': chunk,
                    'chunk_id': i,
                    'source_file': os.path.basename(file_path),
                    'char_count': len(chunk)
                })
                logger.debug(f"分块 {i+1}: {len(chunk)} 字符")
            
            logger.success(f"文档处理完成: {file_path}")
            return results
            
        except Exception as e:
            logger.error(f"处理文档 {file_path} 时出错: {str(e)}")
            raise Exception(f"处理文档 {file_path} 时出错: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """从 PDF 文件提取文本"""
        if PyPDF2 is None:
            logger.error("PyPDF2 库未安装，无法处理 PDF 文件")
            raise ImportError("PyPDF2 库未安装，请运行: pip install PyPDF2")
        
        logger.debug(f"开始提取 PDF 文本: {file_path}")
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                logger.info(f"PDF 文档共 {total_pages} 页")
                
                for i, page in enumerate(pdf_reader.pages):
                    logger.debug(f"正在处理第 {i+1}/{total_pages} 页")
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.debug(f"第 {i+1} 页提取了 {len(page_text)} 字符")
                
                logger.success(f"PDF 文本提取完成，总计 {len(text)} 字符")
                return text
                
        except Exception as e:
            logger.error(f"PDF 文本提取失败: {str(e)}")
            raise Exception(f"PDF 文件读取失败: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """从 DOCX 文件提取文本"""
        if Document is None:
            logger.error("python-docx 库未安装，无法处理 DOCX 文件")
            raise ImportError("python-docx 库未安装，请运行: pip install python-docx")
        
        logger.debug(f"开始提取 DOCX 文本: {file_path}")
        
        try:
            doc = Document(file_path)
            text = ""
            paragraph_count = len(doc.paragraphs)
            logger.info(f"DOCX 文档共 {paragraph_count} 个段落")
            
            for i, paragraph in enumerate(doc.paragraphs):
                if i % 50 == 0:  # 每50个段落记录一次进度
                    logger.debug(f"正在处理段落 {i+1}/{paragraph_count}")
                text += paragraph.text + "\n"
            
            # 提取表格中的文本
            table_count = len(doc.tables)
            if table_count > 0:
                logger.info(f"DOCX 文档共 {table_count} 个表格")
                for i, table in enumerate(doc.tables):
                    logger.debug(f"正在处理表格 {i+1}/{table_count}")
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
            
            logger.success(f"DOCX 文本提取完成，总计 {len(text)} 字符")
            return text
            
        except Exception as e:
            logger.error(f"DOCX 文本提取失败: {str(e)}")
            raise Exception(f"DOCX 文件读取失败: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """从 TXT 文件提取文本"""
        logger.debug(f"开始提取 TXT 文本: {file_path}")
        
        try:
            # 检测文件编码
            logger.debug("正在检测文件编码")
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding']
                confidence = encoding_result['confidence']
                logger.info(f"检测到文件编码: {encoding} (置信度: {confidence:.2f})")
            
            # 使用检测到的编码读取文件
            final_encoding = encoding or 'utf-8'
            logger.debug(f"使用编码 {final_encoding} 读取文件")
            
            with open(file_path, 'r', encoding=final_encoding) as file:
                text = file.read()
                logger.success(f"TXT 文本提取完成，总计 {len(text)} 字符")
                return text
                
        except Exception as e:
            logger.error(f"TXT 文本提取失败: {str(e)}")
            raise Exception(f"TXT 文件读取失败: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """清理和预处理文本"""
        logger.debug("开始文本清理处理")
        original_length = len(text)
        
        # 移除多余的空白字符
        logger.debug("移除多余的空白字符")
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符（保留中文、英文、数字和基本标点）
        logger.debug("移除特殊字符，保留中英文、数字和基本标点")
        text = re.sub(r'[^\u4e00-\u9fff\w\s.,;:!?()\[\]{}"\'-]', ' ', text)
        
        # 移除多余的空行
        logger.debug("移除多余的空行")
        text = re.sub(r'\n\s*\n', '\n', text)
        
        final_length = len(text.strip())
        logger.info(f"文本清理完成: {original_length} -> {final_length} 字符")
        
        return text.strip()
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """将文本分割成块"""
        logger.debug(f"开始文本分块，目标大小: {self.chunk_size}, 重叠: {self.chunk_overlap}")
        
        if len(text) <= self.chunk_size:
            logger.info(f"文本长度 {len(text)} 小于分块大小，返回单个分块")
            return [text]
        
        chunks = []
        start = 0
        chunk_count = 0
        
        while start < len(text):
            # 确定当前块的结束位置
            end = start + self.chunk_size
            
            if end >= len(text):
                # 最后一块
                final_chunk = text[start:]
                chunks.append(final_chunk)
                chunk_count += 1
                logger.debug(f"生成最后分块 {chunk_count}: {len(final_chunk)} 字符")
                break
            
            # 尝试在句号、感叹号或问号处分割
            chunk_text = text[start:end]
            
            # 寻找最佳分割点
            split_point = self._find_best_split_point(chunk_text)
            logger.debug(f"在分块 {chunk_count + 1} 中寻找分割点，结果: {split_point}")
            
            if split_point > 0:
                actual_end = start + split_point
                current_chunk = text[start:actual_end]
                chunks.append(current_chunk)
                chunk_count += 1
                logger.debug(f"生成分块 {chunk_count}: {len(current_chunk)} 字符 (位置 {start}-{actual_end})")
                start = actual_end - self.chunk_overlap
            else:
                # 如果找不到好的分割点，就在指定位置强制分割
                current_chunk = text[start:end]
                chunks.append(current_chunk)
                chunk_count += 1
                logger.debug(f"强制分割生成分块 {chunk_count}: {len(current_chunk)} 字符 (位置 {start}-{end})")
                start = end - self.chunk_overlap
            
            # 确保 start 不会倒退
            if start < 0:
                start = 0
                logger.warning("调整起始位置避免负值")
        
        final_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
        logger.success(f"文本分块完成，共生成 {len(final_chunks)} 个分块")
        
        # 记录分块统计信息
        if final_chunks:
            chunk_sizes = [len(chunk) for chunk in final_chunks]
            avg_size = sum(chunk_sizes) / len(chunk_sizes)
            min_size = min(chunk_sizes)
            max_size = max(chunk_sizes)
            logger.info(f"分块统计 - 平均: {avg_size:.0f}, 最小: {min_size}, 最大: {max_size} 字符")
        
        return final_chunks
    
    def _find_best_split_point(self, text: str) -> int:
        """在文本中找到最佳的分割点"""
        logger.debug(f"在 {len(text)} 字符的文本中寻找最佳分割点")
        
        # 优先级：句号 > 换行符 > 其他标点 > 空格
        split_chars = ['.', '。', '\n', '!', '！', '?', '？', ';', '；', ',', '，', ' ']
        min_position = int(len(text) * 0.7)  # 确保分割点在文本的后70%
        
        logger.debug(f"最小分割位置: {min_position} (70% of {len(text)})")
        
        # 从后往前搜索，找到最佳分割点
        for char in split_chars:
            pos = text.rfind(char)
            logger.debug(f"字符 '{char}' 最后出现位置: {pos}")
            
            if pos > min_position:
                logger.debug(f"找到合适的分割点: 位置 {pos + 1}, 字符 '{char}'")
                return pos + 1
        
        logger.warning("未找到合适的分割点，将强制分割")
        return 0  # 没找到合适的分割点
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """获取文档基本信息"""
        file_size = os.path.getsize(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        return {
            'filename': os.path.basename(file_path),
            'file_size': file_size,
            'file_type': file_ext,
            'supported': file_ext in ['.pdf', '.docx', '.txt']
        }